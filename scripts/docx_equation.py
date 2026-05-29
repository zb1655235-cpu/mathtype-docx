#!/usr/bin/env python3
r"""
Insert MathType/LaTeX equations into .docx files as native OMML equations.

Architecture:
  LaTeX --[Equation.DSMT4 COM]--> MathML --[mathml_to_omml]--> OMML
  OMML --[python-docx XML manipulation]--> .docx file

Usage:
  # Insert single equation
  python docx_equation.py --docx paper.docx --latex "E=mc^2" --pos end

  # Insert after specific text
  python docx_equation.py --docx paper.docx --latex "r_{i,t}=..." --pos "after:事件研究方法"

  # Batch convert text formulas
  python docx_equation.py --docx paper.docx --batch

  # Replace specific paragraph with equation
  python docx_equation.py --docx paper.docx --latex "E=mc^2" --replace 42
"""

import argparse
import sys
import io
from lxml import etree
from pathlib import Path

# Add parent dir to path for latex_to_omml import
sys.path.insert(0, str(Path(__file__).parent))
from latex_to_omml import latex_to_omml, OMML_NS

OMML_Q = f'{{{OMML_NS}}}'


def _is_formula_para(para):
    """Detect if a paragraph looks like a formula (centered, contains math notation)."""
    text = para.text.strip()
    if not text:
        return False
    # Check for math-like content
    math_indicators = ['=', '+', '-', '×', '∑', '√', '∫', 'α', 'β', 'σ', 'μ', 'Δ',
                       '^', '_', '{', '}', '\\frac', '\\sum', '\\sqrt', '\\alpha',
                       'r(i,t)', 'CAR', 'AR', 'DHI']
    score = sum(1 for ind in math_indicators if ind in text)
    return score >= 2


def insert_omath_at_end(doc, omml_elem):
    """Insert an OMML equation at the end of the document."""
    para = doc.add_paragraph()
    para.alignment = 1  # CENTER
    para._element.append(omml_elem)
    return para


def insert_omath_after_text(doc, search_text, omml_elem):
    """Insert an OMML equation after the paragraph containing search_text."""
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            # Insert after this paragraph
            new_para = doc.add_paragraph()
            new_para.alignment = 1  # CENTER
            new_para._element.append(omml_elem)

            # Move the new paragraph right after the found paragraph
            para._element.addnext(new_para._element)
            return new_para
    return None


def insert_omath_before_text(doc, search_text, omml_elem):
    """Insert an OMML equation before the paragraph containing search_text."""
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            new_para = doc.add_paragraph()
            new_para.alignment = 1  # CENTER
            new_para._element.append(omml_elem)

            para._element.addprevious(new_para._element)
            return new_para
    return None


def insert_omath_at_index(doc, index, omml_elem):
    """Insert an OMML equation at a specific paragraph index."""
    if index < 0 or index >= len(doc.paragraphs):
        print(f"  [ERROR] Paragraph index {index} out of range (0-{len(doc.paragraphs)-1})")
        return None

    new_para = doc.add_paragraph()
    new_para.alignment = 1  # CENTER
    new_para._element.append(omml_elem)

    target = doc.paragraphs[index]._element
    target.addprevious(new_para._element)
    return new_para


def replace_paragraph_with_omath(doc, index, omml_elem):
    """Replace a paragraph at index with an OMML equation."""
    if index < 0 or index >= len(doc.paragraphs):
        print(f"  [ERROR] Paragraph index {index} out of range")
        return None

    old_para = doc.paragraphs[index]._element
    parent = old_para.getparent()

    new_para = doc.add_paragraph()
    new_para.alignment = 1  # CENTER
    new_para._element.append(omml_elem)

    parent.replace(old_para, new_para._element)
    return new_para


def batch_convert_formulas(doc):
    """Find and convert all text-based formulas to OMML equations."""
    converted = 0
    skipped = 0

    # Work backwards to avoid index shifting
    indices_to_convert = []
    for i, para in enumerate(doc.paragraphs):
        if _is_formula_para(para) and not _has_omath(para):
            indices_to_convert.append(i)

    for i in reversed(indices_to_convert):
        para = doc.paragraphs[i]
        formula_text = para.text.strip()

        # Extract equation number if present
        eq_num = None
        if formula_text.endswith(')') and '(' in formula_text[-10:]:
            last_paren = formula_text.rfind('(')
            potential_num = formula_text[last_paren:].strip('()')
            if potential_num.isdigit():
                eq_num = potential_num
                formula_text = formula_text[:last_paren].strip()

        # Try to clean up the formula for LaTeX
        latex_formula = _text_to_latex_guess(formula_text)

        if latex_formula:
            try:
                omml, source = latex_to_omml(latex_formula)
                replace_paragraph_with_omath(doc, i, omml)
                converted += 1
            except Exception as e:
                print(f"  [SKIP] P{i}: {str(e)[:80]}")
                skipped += 1
        else:
            skipped += 1

    return converted, skipped


def _has_omath(para):
    """Check if paragraph already contains OMML equation."""
    for child in para._element:
        if child.tag == OMML_Q + 'oMath':
            return True
    return False


def _text_to_latex_guess(text):
    """Attempt to convert a text-based formula to LaTeX.

    Handles common text formula patterns:
      r(i,t) = 100 x (...)  →  r_{i,t}=100\times(...)
      mu_hat(i,e) = ...     →  \hat{\mu}_{i,e}=...
      CAR(i,e)(W) = ...     →  CAR_{i,e}(W)=...
      sigma2_hat(i,e) = ... →  \hat{\sigma}^{2}_{i,e}=...
    """
    latex = text

    # Replace common patterns
    replacements = [
        # Function notation with indices: r(i,t) → r_{i,t}
        # (This is heuristic - we handle simple cases)
        (' × ', r' \times '),
        (' − ', ' - '),
        (' ≤ ', r' \leq '),
        (' ≥ ', r' \geq '),
        (' √(', r' \sqrt{'),
        (' Σ_{', r' \sum_{'),
        (' ⋯ ', r' \cdots '),
        (' α ', r' \alpha '),
        (' α_', r' \alpha_'),
        (' σ̂', r' \hat{\sigma}'),
        (' μ̂', r' \hat{\mu}'),
        (' Δy', r' \Delta y'),
        (' T₀', ' T_0'),
        (' L_W', ' L_W'),
        (' τ₁', r' \tau_1'),
        (' τ₂', r' \tau_2'),
        (' α_Bonf', r' \alpha_{Bonf}'),
        (' α / M', r' \alpha / M'),
        (' p₍₁₎', ' p_{(1)}'),
        (' p₍₂₎', ' p_{(2)}'),
        (' p₍ₖ₎', ' p_{(k)}'),
        (' p(M)', ' p_{(M)}'),
        (' DHI_gold', ' DHI^{gold}'),
        (' CAR_gold', ' CAR^{gold}'),
    ]

    for old, new in replacements:
        latex = latex.replace(old, new)

    # Handle function(index,index) → function_{index,index}
    # This is a complex heuristic, so we handle specific known cases
    known_funcs = ['r', 'P', 'AR', 'CAR', 'ln']
    for func in known_funcs:
        # Find func(...) but not already with braces
        import re
        pattern = re.escape(func) + r'\(([^)]+)\)'
        latex = re.sub(pattern, func + r'_{\1}', latex)

    return latex


def main():
    parser = argparse.ArgumentParser(
        description='Insert MathType/LaTeX equations into .docx files as OMML'
    )
    parser.add_argument('--docx', type=str, required=True, help='Path to .docx file')
    parser.add_argument('--latex', type=str, help='LaTeX formula to insert')
    parser.add_argument('--pos', type=str, default='end',
                       help='Insert position: "end", "after:TEXT", "before:TEXT", or paragraph index')
    parser.add_argument('--replace', type=int, help='Replace paragraph at index with equation')
    parser.add_argument('--batch', action='store_true',
                       help='Batch convert all text formulas in the document')
    parser.add_argument('--eqnum', type=str, help='Equation number (appended after formula)')
    parser.add_argument('--output', type=str, help='Output path (defaults to overwrite input)')

    args = parser.parse_args()

    from docx import Document

    doc = Document(args.docx)
    print(f"Opened: {args.docx}")
    print(f"  Paragraphs: {len(doc.paragraphs)}")

    if args.batch:
        print("\n=== Batch converting text formulas to OMML ===")
        converted, skipped = batch_convert_formulas(doc)
        print(f"\n  Converted: {converted}")
        print(f"  Skipped: {skipped}")

    elif args.latex:
        latex_input = args.latex
        if args.eqnum:
            latex_input += f' \\qquad ({args.eqnum})'

        print(f"\n=== Converting LaTeX to OMML ===")
        print(f"  Input: {latex_input}")

        omml, source = latex_to_omml(latex_input)
        print(f"  Source: {source}")
        print(f"  OMML: {etree.tostring(omml, encoding='unicode')[:200]}...")

        # Insert into document
        if args.replace is not None:
            para = replace_paragraph_with_omath(doc, args.replace, omml)
            if para:
                print(f"  Replaced paragraph {args.replace} with equation")
        elif args.pos == 'end':
            para = insert_omath_at_end(doc, omml)
            print(f"  Inserted at end of document")
        elif args.pos.startswith('after:'):
            search = args.pos[6:]
            para = insert_omath_after_text(doc, search, omml)
            if para:
                print(f"  Inserted after text: '{search}'")
            else:
                print(f"  [ERROR] Text not found: '{search}'")
        elif args.pos.startswith('before:'):
            search = args.pos[7:]
            para = insert_omath_before_text(doc, search, omml)
            if para:
                print(f"  Inserted before text: '{search}'")
            else:
                print(f"  [ERROR] Text not found: '{search}'")
        elif args.pos.isdigit():
            para = insert_omath_at_index(doc, int(args.pos), omml)
            if para:
                print(f"  Inserted at paragraph index {args.pos}")
        else:
            print(f"  [ERROR] Invalid position: '{args.pos}'")
            return 1

    else:
        parser.print_help()
        return 1

    # Save
    output_path = args.output if args.output else args.docx
    doc.save(output_path)
    print(f"\nSaved: {output_path}")
    return 0


if __name__ == '__main__':
    sys.exit(main())
